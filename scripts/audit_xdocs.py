"""Cross-document consistency audit: main.tex vs RESULTS.md vs README.md vs paper.docx"""
import docx, json, re, sys
from pathlib import Path

BASE = Path(__file__).resolve().parent.parent

def load_docx(path):
    doc = docx.Document(str(path))
    text = '\n'.join(p.text for p in doc.paragraphs)
    tables = []
    for t in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        tables.append(rows)
    return text, tables

def load_file(path):
    return Path(path).read_text()

# Load all 4 documents
tex = load_file(BASE / 'paper' / 'main.tex')
rmd = load_file(BASE / 'RESULTS.md')
readme = load_file(BASE / 'README.md')
_docx_matches = list(BASE.glob('*.docx'))
if _docx_matches:
    docx_text, docx_tables = load_docx(_docx_matches[0])
else:
    docx_text, docx_tables = '', []

issues = []

def check(label, tex_val, rmd_val, readme_val, docx_val,
           tex_found=True, rmd_found=True, readme_found=True, docx_found=True):
    """Compare value across all 4 documents."""
    if tex_found and rmd_found:
        if tex_val != rmd_val:
            issues.append(f"MISMATCH [{label}]: tex={tex_val} vs rmd={rmd_val}")
    if tex_found and readme_found:
        if tex_val != readme_val:
            issues.append(f"MISMATCH [{label}]: tex={tex_val} vs readme={readme_val}")
    if tex_found and docx_found:
        if tex_val != docx_val:
            issues.append(f"MISMATCH [{label}]: tex={tex_val} vs docx={docx_val}")
    if not tex_found:
        issues.append(f"MISSING [tex] {label}")
    if not rmd_found:
        issues.append(f"MISSING [rmd] {label}")
    if not readme_found:
        issues.append(f"MISSING [readme] {label}")
    if not docx_found:
        issues.append(f"MISSING [docx] {label}")

def extract_tex_value(pattern, text, group=1):
    m = re.search(pattern, text)
    if m is None:
        return None
    try:
        return m.group(group)
    except IndexError:
        return m.group(0)

def extract_rmd_value(pattern):
    matches = re.findall(pattern, rmd)
    return matches

# =====================================================
# TABLE 1: Training Performance
# =====================================================
tex_train_acc_clean = extract_tex_value(r'Train Accuracy.*?96\.68%\\[[^\\]+\\]', tex)
tex_train_acc_corr = extract_tex_value(r'Train Accuracy.*?94\.21%\\[[^\\]+\\]', tex)
tex_test_acc_clean = extract_tex_value(r'Test Accuracy.*?95\.31%\\[[^\\]+\\]', tex)
tex_test_acc_corr = extract_tex_value(r'Test Accuracy.*?93\.71%\\[[^\\]+\\]', tex)

rmd_train_acc_clean = extract_tex_value(r'Train Accuracy.*?96\.68%', rmd)
rmd_train_acc_corr = extract_tex_value(r'Train Accuracy.*?94\.21%', rmd)
rmd_test_acc_clean = extract_tex_value(r'Test Accuracy.*?95\.31%', rmd)
rmd_test_acc_corr = extract_tex_value(r'Test Accuracy.*?93\.71%', rmd)

readme_train_acc_clean = extract_tex_value(r'Train Accuracy.*?96\.68%', readme)
readme_train_acc_corr = extract_tex_value(r'Train Accuracy.*?94\.21%', readme)
readme_test_acc_clean = extract_tex_value(r'Test Accuracy.*?95\.31%', readme)
readme_test_acc_corr = extract_tex_value(r'Test Accuracy.*?93\.71%', readme)

check('Train Acc Clean', tex_train_acc_clean, rmd_train_acc_clean, readme_train_acc_clean, None)
check('Train Acc Corrupted', tex_train_acc_corr, rmd_train_acc_corr, readme_train_acc_corr, None)
check('Test Acc Clean', tex_test_acc_clean, rmd_test_acc_clean, readme_test_acc_clean, None)
check('Test Acc Corrupted', tex_test_acc_corr, rmd_test_acc_corr, readme_test_acc_corr, None)

# =====================================================
# TABLE 2: Weight Geometry (MNIST)
# =====================================================
pairs_phase1 = [
    ('FC1 Spectral Norm Clean', r'FC1 Spectral Norm.*?4\.37'),
    ('FC1 Spectral Norm Corr', r'FC1 Spectral Norm.*?3\.66'),
    ('FC2 Spectral Norm Clean', r'FC2 Spectral Norm.*?2\.58'),
    ('FC2 Spectral Norm Corr', r'FC2 Spectral Norm.*?1\.39'),
    ('FC1 Frobenius Clean', r'FC1 Frobenius Norm.*?12\.56'),
    ('FC1 Frobenius Corr', r'FC1 Frobenius Norm.*?10\.33'),
    ('FC2 Frobenius Clean', r'FC2 Frobenius Norm.*?4\.99'),
    ('FC2 Frobenius Corr', r'FC2 Frobenius Norm.*?2\.79'),
]

for label, pat in pairs_phase1:
    tex_v = extract_tex_value(pat, tex)
    rmd_v = extract_tex_value(pat, rmd)
    readme_v = extract_tex_value(pat, readme)
    check(label, tex_v, rmd_v, readme_v, None)

# =====================================================
# TABLE 3: CKA (MNIST)
# =====================================================
pairs_cka = [
    ('CKA fc1_pre->fc1_post Clean', r'CKA.*?fc1_pre.*?fc1_post.*?0\.850'),
    ('CKA fc1_pre->fc1_post Corr', r'CKA.*?fc1_pre.*?fc1_post.*?0\.690'),
    ('CKA fc1_pre->fc1_post Delta', r'\+0\.160'),
]

for label, pat in pairs_cka:
    m_tex = re.search(pat, tex)
    m_rmd = re.search(pat, rmd)
    check(label,
          m_tex.group() if m_tex else None,
          m_rmd.group() if m_rmd else None,
          re.search(pat, readme).group() if re.search(pat, readme) else None,
          None)

# =====================================================
# TABLE 4: ROME MNIST
# =====================================================
pairs_rome = [
    ('ROME fc1 avg Clean', r'14\.49'),
    ('ROME fc1 avg Corrupted', r'7\.18'),
    ('ROME fc2 avg Clean', r'19\.08'),
    ('ROME fc2 avg Corrupted', r'4\.40'),
    ('ROME fc1 Ratio', r'2\.02'),
    ('ROME fc2 Ratio', r'4\.34'),
]

for label, pat in pairs_rome:
    m_tex = re.search(pat, tex)
    m_rmd = re.search(pat, rmd)
    check(label,
          m_tex.group() if m_tex else None,
          m_rmd.group() if m_rmd else None,
          re.search(pat, readme).group() if re.search(pat, readme) else None,
          None)

# =====================================================
# TABLE 5: ROME CIFAR-10
# =====================================================
pairs_rome_cifar = [
    ('CIFAR-10 ROME ratio', r'1\.84'),
    ('CIFAR-10 ROME p<0.05', r'p\s*<\s*0\.05'),
]

for label, pat in pairs_rome_cifar:
    m_tex = re.search(pat, tex)
    m_rmd = re.search(pat, rmd)
    m_readme = re.search(pat, readme)
    check(label,
          m_tex.group() if m_tex else None,
          m_rmd.group() if m_rmd else None,
          m_readme.group() if m_readme else None,
          None)

# =====================================================
# TABLE 6: Phase 3
# =====================================================
pairs_phase3 = [
    ('Phase3 Loss Gap Clean', r'\+0\.442'),
    ('Phase3 Loss Gap Corrupted', r'\-0\.051'),
]

for label, pat in pairs_phase3:
    m_tex = re.search(pat, tex)
    m_rmd = re.search(pat, rmd)
    check(label,
          m_tex.group() if m_tex else None,
          m_rmd.group() if m_rmd else None,
          re.search(pat, readme).group() if re.search(pat, readme) else None,
          None)

# =====================================================
# TABLE 7: Random baseline
# =====================================================
pairs_rb = [
    ('Random Baseline 0%', r'0\.00%'),
    ('Signal Ratio infin', r'∞'),
]

for label, pat in pairs_rb:
    m_rmd = re.search(pat, rmd)
    check(label,
          re.search(pat, tex).group() if re.search(pat, tex) else None,
          m_rmd.group() if m_rmd else None,
          re.search(pat, readme).group() if re.search(pat, readme) else None,
          None)

# =====================================================
# TABLE 8: Gradient alignment
# =====================================================
g_tex = re.search(r'\+0\.9944\s*\[0\.9926,\s*0\.9961\]', tex)
g_rmd = re.search(r'\+0\.9944\s*\[0\.9926,\s*0\.9961\]', rmd)
g_readme = re.search(r'\+0\.9944\s*\[0\.9926,\s*0\.9961\]', readme)
check('Gradient Align +0.9944',
      g_tex.group() if g_tex else None,
      g_rmd.group() if g_rmd else None,
      g_readme.group() if g_readme else None, None)

# =====================================================
# TABLE 9: Multi-class ROME recovery values
# =====================================================
mc_values = [
    ('7->1 Recovery', r'\+0\.141'),
    ('1->7 Recovery', r'\+0\.217'),
    ('5->6 Recovery', r'\+0\.120'),
    ('0->8 Recovery', r'\+0\.097'),
    ('Multi-layer ROME 22-40%', r'22%'),
]
for label, pat in mc_values:
    m_tex = re.search(pat, tex)
    m_rmd = re.search(pat, rmd)
    m_readme = re.search(pat, readme)
    check(label,
          m_tex.group() if m_tex else None,
          m_rmd.group() if m_rmd else None,
          m_readme.group() if m_readme else None, None)

# =====================================================
# Check RESULTS.md rank ablation matches paper
# =====================================================
rank_pairs = [
    ('Rank 5 Clean', r'68\.4%'),
    ('Rank 5 Corrupted', r'62\.3%'),
    ('Rank Rank-5 Gap 6.1pp', r'6\.1'),
]
for label, pat in rank_pairs:
    m_rmd = re.search(pat, rmd)
    m_readme = re.search(pat, readme)
    check(label,
          None,  # not in paper in detail
          m_rmd.group() if m_rmd else None,
          m_readme.group() if m_readme else None, None)

# =====================================================
# DOCX body text checks
# =====================================================
docx_checks = [
    ("ROME ratio mention", r"1\.84"),
    ("gradient alignment mention", r"0\.9944"),
    ("random baseline mention", r"0%"),
    ("multi-layer ROME mention", r"multi.lay"),
    ("CIFAR-10 CKA mention", r"0\.103"),
    ("circuit size 0", r"0\b"),
]
for label, pat in docx_checks:
    m = re.search(pat, docx_text, re.IGNORECASE)
    if m:
        issues.append(f"OK [docx] {label}: found '{m.group()}'")
    else:
        issues.append(f"MISSING [docx] {label}: pattern '{pat}' not found in body")

# =====================================================
# DOCX tables check (spot key values)
# =====================================================
docx_table_text = '\n'.join(
    ' | '.join(row) for table in docx_tables for row in table
)

docx_table_checks = [
    ("DOCX table: 4.37 spectral norm", r"4\.37"),
    ("DOCX table: 2.58 spectral norm", r"2\.58"),
    ("DOCX table: 12.56 Frobenius", r"12\.56"),
    ("DOCX table: train acc 96.68%", r"96\.68"),
    ("DOCX table: train acc 94.21%", r"94\.21"),
    ("DOCX table: test acc 95.31%", r"95\.31"),
    ("DOCX table: test acc 93.71%", r"93\.71"),
    ("DOCX table: CKA delta 0.160", r"0\.160"),
    ("DOCX table: CKA delta 0.103", r"0\.103"),
    ("DOCX table: ROME ratio 2.02", r"2\.02"),
    ("DOCX table: ROME ratio 4.34", r"4\.34"),
    ("DOCX table: ROME ratio 1.84", r"1\.84"),
    ("DOCX table: rank 5/10 68.4%", r"68\.4"),
    ("DOCX table: rank 5/10 62.3%", r"62\.3"),
    ("DOCX table: 0.9944 gradient", r"0\.9944"),
    ("DOCX table: random baseline 0%", r"0\.00%"),
    ("DOCX table: 7->1 recovery", r"\+0\.141"),
    ("DOCX table: 1->7 recovery", r"\+0\.217"),
    ("DOCX table: 5->6 recovery", r"\+0\.120"),
    ("DOCX table: 0->8 recovery", r"\+0\.097"),
    ("DOCX table: multi-layer 22-40%", r"22"),
]
for label, pat in docx_table_checks:
    m = re.search(pat, docx_table_text)
    if m:
        issues.append(f"OK [docx] {label}: found '{m.group()}'")
    else:
        issues.append(f"MISSING [docx] {label}: pattern '{pat}' not found in tables")

# =====================================================
# RESULTS.md section completness check
# =====================================================
rmd_sections = [
    "Training Performance",
    "Phase 1: Weight Geometry",
    "Phase 2: Representation Similarity",
    "Phase 3: Influence & Memorization",
    "Phase 4: ROME",
    "Multi-Class ROME",
    "Random Baseline for ROME",
    "Gradient Anti-Alignment",
    "Rank Ablation",
    "Scaling Analysis",
    "CIFAR-10",
]
for sec in rmd_sections:
    if sec.lower() in rmd.lower():
        issues.append(f"OK [rmd] Section '{sec}' present")
    else:
        issues.append(f"MISSING [rmd] Section '{sec}'")

# =====================================================
# README key numbers check
# =====================================================
readme_key_values = [
    ("README: train acc clean", r"96\.68%"),
    ("README: train acc corr", r"94\.21%"),
    ("README: test acc clean", r"95\.31%"),
    ("README: test acc corr", r"93\.71%"),
    ("README: CKA delta 0.160", r"\+0\.160"),
    ("README: ROME fc1 ratio 2.02", r"2\.02"),
    ("README: ROME fc2 ratio 4.34", r"4\.34"),
    ("README: CIFAR-10 ratio 1.84", r"1\.84"),
    ("README: gradient align", r"0\.9944"),
    ("README: loss gap", r"0\.051"),
    ("README: rank-5 gap 6.1", r"6\.1"),
    ("README: random baseline", r"∞"),
    ("README: multi-layer 22-40%", r"22"),
]
for label, pat in readme_key_values:
    m = re.search(pat, readme)
    if m:
        issues.append(f"OK [readme] {label}: found '{m.group()}'")
    else:
        issues.append(f"MISSING [readme] {label}: pattern '{pat}' not found")

# =====================================================
# Print results
# =====================================================
print("=" * 72)
print("CROSS-DOCUMENT AUDIT RESULTS")
print("=" * 72)

errors = [i for i in issues if i.startswith('MISMATCH')]
missing = [i for i in issues if i.startswith('MISSING')]
oks = [i for i in issues if i.startswith('OK')]

if errors:
    print(f"\nERRORS ({len(errors)}):")
    for e in errors:
        print(f"  ✗ {e}")

if missing:
    print(f"\nMISSING ({len(missing)}):")
    for m in missing:
        print(f"  ? {m}")

print(f"\nOK ({len(oks)}):")
for o in oks:
    print(f"  ✓ {o}")

print(f"\n{'─' * 72}")
print(f"Summary: {len(errors)} errors, {len(missing)} missing, {len(oks)} ok")
print(f"{'─' * 72}")

sys.exit(1 if errors else 0)
